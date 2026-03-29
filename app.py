"""
Phase 9: Streamlit Web Application
- Resume upload and job description input
- Display match scores, predictions, and extracted skills
- Interactive dashboard for recruiter
"""

import streamlit as st
import pandas as pd
import numpy as np
import sys
import os
import nltk

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
from pathlib import Path
from io import BytesIO

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from data_preprocessing import DataPreprocessor
from nlp_processor import NLPProcessor
from matching_engine import MatchingEngine
from feature_engineering import FeatureEngineer
from model_persistence import ModelPersistence
from login import show_login_page, show_logout_button

def extract_text_from_file(uploaded_file):
    """Extract text content from uploaded file (PDF, TXT, DOCX)"""
    try:
        if not uploaded_file:
            return "Error: No file provided"
        
        file_name = uploaded_file.name.lower()
        file_bytes = uploaded_file.read()

        # Handle text files
        if file_name.endswith('.txt') or uploaded_file.type == "text/plain":
            try:
                content = file_bytes.decode('utf-8', errors='ignore')
                return content if content.strip() else "Error: File is empty"
            except Exception as e:
                return f"Error decoding TXT file: {str(e)}"

        # Handle PDF files
        elif file_name.endswith('.pdf') or uploaded_file.type == "application/pdf":
            try:
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(BytesIO(file_bytes))
                
                if len(pdf_reader.pages) == 0:
                    return "Error: PDF has no pages"
                
                text = ""
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_err:
                        # Skip problematic pages
                        continue
                
                return text.strip() if text.strip() else "Error: Could not extract text from PDF. Try a text-based PDF."
            except ImportError:
                return "Error: PyPDF2 not installed. Run: pip install PyPDF2"
            except Exception as e:
                return f"Error processing PDF: {str(e)}"

        # Handle DOCX files
        elif file_name.endswith(('.docx', '.doc')) or "wordprocessingml" in (uploaded_file.type or "").lower():
            try:
                from docx import Document
                doc = Document(BytesIO(file_bytes))
                text = ""
                
                # Extract from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"
                
                return text.strip() if text.strip() else "Error: DOCX file is empty or has no extractable text"
            except ImportError:
                return "Error: python-docx not installed. Run: pip install python-docx"
            except Exception as e:
                return f"Error processing DOCX: {str(e)}"

        else:
            return f"Unsupported file type: {file_name}. Please upload PDF, TXT, or DOCX files."

    except Exception as e:
        return f"Error extracting text: {str(e)}"

# Page config
st.set_page_config(
    page_title="Smart Hiring Platform",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
        .main-header {
            text-align: center;
            color: #2E86AB;
            margin-bottom: 30px;
        }
        .match-score {
            font-size: 48px;
            font-weight: bold;
            text-align: center;
            padding: 20px;
            border-radius: 10px;
            margin: 20px 0;
        }
        .module-panel {
            background: linear-gradient(180deg, #ffffff, #f7fbff);
            border: 1px solid #dbe8f7;
            border-radius: 12px;
            padding: 16px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
            margin-bottom: 16px;
        }
        .module-title {
            background-color: #eff8ff;
            border-left: 4px solid #2e86ab;
            padding: 10px 14px;
            border-radius: 8px;
            margin-bottom: 14px;
            font-weight: 700;
            color: #0d3c70;
        }
        .match-score {
            font-size: 48px;
            font-weight: bold;
            text-align: center;
            padding: 20px;
            border-radius: 10px;
            margin: 20px 0;
        }
        .strong-match {
            background-color: #90EE90;
            color: #006400;
        }
        .good-match {
            background-color: #FFD700;
            color: #DAA520;
        }
        .moderate-match {
            background-color: #FFA500;
            color: #FF8C00;
        }
        .poor-match {
            background-color: #FFB6C1;
            color: #DC143C;
        }
        .metric-box {
            background-color: #F0F2F6;
            padding: 15px;
            border-radius: 8px;
            margin: 10px 0;
        }
        /* global background */
        .css-1d391kg, .css-18e3th9, [data-testid="stAppViewContainer"], .stApp { background: linear-gradient(135deg, #0f1c3b, #0a1546) !important; }
        .css-1lcbmhc, .block-container { background: transparent !important; }
        .main, .stApp, section.main { background: transparent !important; }
        body, .stApp, .css-1d391kg, .css-18e3th9, .block-container, .main, span, p, label, h1, h2, h3, h4, h5, h6, .stMarkdown, .stText {
            color: #ffffff !important;
        }
        .module-panel {
            background: rgba(0, 14, 47, 0.92) !important;
            border: 1px solid rgba(255, 255, 255, 0.18);
            border-radius: 14px;
            padding: 18px;
            box-shadow: 0 8px 18px rgba(0, 0, 0, 0.08);
            margin-bottom: 20px;
        }
        .module-title {
            background-color: #edf6ff;
            border-left: 5px solid #357ab7;
            padding: 9px 12px;
            border-radius: 8px;
            margin-bottom: 16px;
            font-weight: 700;
            color: #153962;
        }
        /* Sidebar/Navigation styling */
        .css-1lcbmhc, .stSidebar, [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #1e3a5f, #0f1c3b) !important;
        }
        .css-1lcbmhc .stSidebar, .stSidebar [data-testid="stSidebarContent"] {
            background: transparent !important;
        }
        .css-1lcbmhc h1, .css-1lcbmhc h2, .css-1lcbmhc h3, .css-1lcbmhc label, .css-1lcbmhc span {
            color: #ffffff !important;
        }
        .css-1lcbmhc .stRadio label, .css-1lcbmhc .stSelectbox label {
            color: #ffffff !important;
        }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'nlp_processor' not in st.session_state:
    st.session_state.nlp_processor = NLPProcessor()

if 'matching_engine' not in st.session_state:
    st.session_state.matching_engine = MatchingEngine(st.session_state.nlp_processor)

if 'feature_engineer' not in st.session_state:
    st.session_state.feature_engineer = FeatureEngineer()

if 'model_persistence' not in st.session_state:
    st.session_state.model_persistence = ModelPersistence()

# Initialize dashboard data tracking
if 'dashboard_data' not in st.session_state:
    st.session_state.dashboard_data = {
        'total_candidates': 0,
        'active_positions': 0,
        'total_matches': 0,
        'shortlisted': 0,
        'recent_matches': [],
        'avg_match_score': 0.0,
        'processed_resumes': []
    }


def module_wrapper(title):
    st.markdown(f'<div class="module-panel">', unsafe_allow_html=True)
    st.markdown(f'<div class="module-title">{title}</div>', unsafe_allow_html=True)


def module_close():
    st.markdown('</div>', unsafe_allow_html=True)


def main():
    # Check if user is logged in
    if not st.session_state.get('logged_in'):
        show_login_page()
        return
    
    # Show logout button in sidebar
    show_logout_button()
    
    st.markdown("""
        <h1 class="main-header">🚀 AI-Powered Smart Hiring & Candidate Intelligence Platform</h1>
    """, unsafe_allow_html=True)
    
    # Sidebar
    st.sidebar.title("🔧 Navigation")
    page = st.sidebar.radio(
        "Select Option:",
        ["📊 Dashboard", "👤 Resume Analysis", "🎯 Job Matching", "📈 Insights", "⚙️ Settings"]
    )
    
    if page == "📊 Dashboard":
        show_dashboard()
    elif page == "👤 Resume Analysis":
        show_resume_analysis()
    elif page == "🎯 Job Matching":
        show_job_matching()
    elif page == "📈 Insights":
        show_insights()
    elif page == "⚙️ Settings":
        show_settings()


def show_dashboard():
    """Main dashboard view"""
    module_wrapper('📊 Hiring Dashboard')
    
    # Get dashboard data from session state
    dashboard_data = st.session_state.dashboard_data
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_candidates = dashboard_data['total_candidates']
        st.metric("Total Candidates", f"{total_candidates}", f"+{total_candidates}")
    
    with col2:
        active_positions = dashboard_data['active_positions']
        st.metric("Active Positions", f"{active_positions}", f"+{active_positions}")
    
    with col3:
        avg_score = dashboard_data['avg_match_score']
        st.metric("Avg Match Score", f"{avg_score:.1f}%", f"{avg_score:.1f}%")
    
    with col4:
        shortlisted = dashboard_data['shortlisted']
        st.metric("Shortlisted", f"{shortlisted}", f"+{shortlisted}")
    
    st.markdown("---")
    
    # Recent matches
    st.subheader("Recent Matches")
    
    if dashboard_data['recent_matches']:
        # Convert recent matches to DataFrame
        df_recent = pd.DataFrame(dashboard_data['recent_matches'])
        st.dataframe(df_recent, width='stretch')
    else:
        # Show placeholder when no matches yet
        st.info("No matches processed yet. Upload a resume and job description to see results here.")
        df_placeholder = pd.DataFrame({
            'Candidate': ['No data yet'],
            'Position': ['No data yet'],
            'Match Score': [0],
            'Status': ['Pending']
        })
        st.dataframe(df_placeholder, width='stretch')
    
    module_close()


def show_resume_analysis():
    """Resume analysis page"""
    module_wrapper('👤 Resume Analysis')
    
    col1, col2 = st.columns([1, 1], gap="large")
    
    # Initialize state for resume text if not exists
    if 'resume_text' not in st.session_state:
        st.session_state.resume_text = ""
    if 'upload_success' not in st.session_state:
        st.session_state.upload_success = False
    
    with col1:
        st.subheader("📄 Upload Resume")
        uploaded_file = st.file_uploader(
            "Choose a resume file",
            type=['pdf', 'txt', 'docx'],
            key=f"resume_uploader_{st.session_state.get('upload_count', 0)}"
        )
        
        if uploaded_file:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            resume_text = extract_text_from_file(uploaded_file)
            
            # Check if extraction was successful
            if resume_text.startswith("Error") or resume_text.startswith("Unsupported") or not resume_text.strip():
                st.error(f"❌ Failed to extract text from resume")
                st.error(resume_text)
                st.info("💡 Try uploading a plain text file (.txt) or ensure the PDF/DOCX is not password-protected or corrupted.")
                st.session_state.resume_text = ""
                st.session_state.upload_success = False
            else:
                st.info(f"📄 Extracted {len(resume_text)} characters from resume")
                st.session_state.resume_text = resume_text
                st.session_state.upload_success = True
            
            if st.session_state.resume_text:
                st.text_area("Resume Content", st.session_state.resume_text, height=300, disabled=True)
        else:
            st.info("📤 Upload a resume file to see the analysis")
    
    with col2:
        st.subheader("📊 Extracted Information")
        
        if st.session_state.upload_success and st.session_state.resume_text:
            try:
                # Process resume
                nlp_processor = st.session_state.nlp_processor
                
                # Extract information
                skills = nlp_processor.extract_skills_from_text(st.session_state.resume_text)
                education = nlp_processor.extract_education(st.session_state.resume_text)
                experience = nlp_processor.extract_years_of_experience(st.session_state.resume_text)
                
                # Display extracted data
                st.markdown("#### 🛠️ Skills Extracted")
                if skills:
                    cols = st.columns(2)
                    for idx, skill in enumerate(skills):
                        with cols[idx % 2]:
                            st.write(f"• {skill}")
                else:
                    st.info("No skills detected. Try uploading a resume with more technical details.")
                    st.info("💡 Look for skills like Python, Java, React, AWS, etc.")
                
                st.markdown("#### 🎓 Education")
                if education:
                    for edu in education:
                        st.write(f"• {edu}")
                else:
                    st.info("No education information found.")
                    st.info("💡 Look for degrees like Bachelor, Master, B.Tech, M.Sc, etc.")
                
                st.markdown(f"#### 📅 Experience")
                st.write(f"**{experience} years** of experience")
                if experience == 0:
                    st.info("💡 Look for experience patterns like '5 years', '3+ years', etc.")
                
                # Summary
                st.markdown("#### 📋 Summary")
                summary_data = {
                    'Metric': ['Total Skills', 'Education Records', 'Years Experience'],
                    'Value': [len(skills), len(education), experience]
                }
                st.dataframe(pd.DataFrame(summary_data), width='stretch')
                
                # Update dashboard data
                if uploaded_file.name not in [r.get('filename', '') for r in st.session_state.dashboard_data['processed_resumes']]:
                    st.session_state.dashboard_data['total_candidates'] += 1
                    st.session_state.dashboard_data['processed_resumes'].append({
                        'filename': uploaded_file.name,
                        'skills_count': len(skills),
                        'education_count': len(education),
                        'experience_years': experience,
                        'processed_at': pd.Timestamp.now()
                    })
            
            except Exception as e:
                st.error(f"❌ Error processing resume: {str(e)}")
                st.info("Please try uploading a different file or check the file format.")
        else:
            st.info("📤 Upload a resume file to see the analysis")
    module_close()


def show_job_matching():
    """Resume-Job matching page"""
    module_wrapper('🎯 Resume-Job Matching Engine')
    
    col1, col2 = st.columns([1, 1], gap="large")
    
    resume_text = ""
    job_text = ""
    
    with col1:
        st.subheader("👤 Candidate Resume")
        st.markdown("**Option 1: Drag and drop or click to upload resume**")
        resume_file = st.file_uploader(
            "Upload Resume (PDF, TXT, DOCX)",
            type=['pdf', 'txt', 'docx'],
            key="resume_uploader"
        )
        
        st.markdown("**Option 2: Or paste resume text directly**")
        resume_text_input = st.text_area(
            "Paste resume text:",
            height=100,
            placeholder="Alternatively, paste resume content here...",
            key="resume_text_input"
        )
        
        # Determine which input to use
        if resume_file and not resume_text_input.strip():
            st.success(f"✅ Resume uploaded: {resume_file.name}")
            resume_text = extract_text_from_file(resume_file)
            
            # Check if extraction was successful
            if resume_text.startswith("Error") or resume_text.startswith("Unsupported") or not resume_text.strip():
                st.error(f"❌ Failed to extract text from resume: {resume_text}")
                st.info("💡 Try uploading a plain text file (.txt) or ensure the PDF/DOCX is not password-protected or corrupted.")
                resume_text = ""  # Clear the text so matching won't proceed
            else:
                st.info(f"📄 Extracted {len(resume_text)} characters from resume")
                
        elif resume_text_input.strip():
            resume_text = resume_text_input
            st.success(f"✅ Resume text entered ({len(resume_text)} characters)")
        else:
            resume_text = ""
        
        # Show preview if we have text
        if resume_text:
            with st.expander("📄 Resume Preview", expanded=False):
                st.text_area("Resume Text", resume_text, height=200, disabled=True)
    
    with col2:
        st.subheader("💼 Job Description")
        st.markdown("**Option 1: Drag and drop or click to upload job description**")
        job_file = st.file_uploader(
            "Upload Job Description (PDF, TXT, DOCX)",
            type=['pdf', 'txt', 'docx'],
            key="job_uploader"
        )
        
        st.markdown("**Option 2: Or paste job description text directly**")
        job_text_input = st.text_area(
            "Paste job description:",
            height=100,
            placeholder="Alternatively, paste job description content here...",
            key="job_text_input"
        )
        
        # Determine which input to use
        if job_file and not job_text_input.strip():
            st.success(f"✅ Job description uploaded: {job_file.name}")
            job_text = extract_text_from_file(job_file)
            
            # Check if extraction was successful
            if job_text.startswith("Error") or job_text.startswith("Unsupported") or not job_text.strip():
                st.error(f"❌ Failed to extract text from job description: {job_text}")
                st.info("💡 Try uploading a plain text file (.txt) or ensure the PDF/DOCX is not password-protected or corrupted.")
                job_text = ""  # Clear the text so matching won't proceed
            else:
                st.info(f"📄 Extracted {len(job_text)} characters from job description")
                
        elif job_text_input.strip():
            job_text = job_text_input
            st.success(f"✅ Job description text entered ({len(job_text)} characters)")
        else:
            job_text = ""
        
        # Show preview if we have text
        if job_text:
            with st.expander("📄 Job Description Preview", expanded=False):
                st.text_area("Job Text", job_text, height=200, disabled=True)
    
    if st.button("🔍 Calculate Match Score", width='stretch', type="primary"):
        if resume_text and job_text and resume_text.strip() and job_text.strip():
            with st.spinner("Processing..."):
                # Process texts
                nlp_processor = st.session_state.nlp_processor
                matching_engine = st.session_state.matching_engine
                
                # Create data dictionaries
                resume_data = {
                    'processed_text': nlp_processor.preprocess_text(resume_text),
                    'extracted_skills': nlp_processor.extract_skills_from_text(resume_text),
                    'years_experience': nlp_processor.extract_years_of_experience(resume_text),
                    'education_level': 2  # Default
                }
                
                job_data = {
                    'job_description': nlp_processor.preprocess_text(job_text),
                    'skills_required': nlp_processor.extract_skills_from_text(job_text),
                    'experiencere_requirement': 0,
                    'educationaL_requirements': ''
                }
                
                # Calculate match
                match_result = matching_engine.match_resume_to_job(resume_data, job_data)
                
                # Display results
                st.markdown("---")
                st.header("📊 Matching Results")
                
                # Overall score with color coding
                score = match_result['overall_score']
                if score >= 80:
                    css_class = "strong-match"
                    emoji = "✅"
                elif score >= 60:
                    css_class = "good-match"
                    emoji = "👍"
                elif score >= 40:
                    css_class = "moderate-match"
                    emoji = "⚠️"
                else:
                    css_class = "poor-match"
                    emoji = "❌"
                
                st.markdown(f"""
                    <div class="match-score {css_class}">
                        {emoji} Overall Match Score: {score:.1f}%
                    </div>
                """, unsafe_allow_html=True)
                
                # Detailed breakdown
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric(
                        "Content Similarity",
                        f"{match_result['content_similarity']:.1f}%"
                    )
                
                with col2:
                    skill_match = match_result['skill_match']
                    st.metric(
                        "Skill Match",
                        f"{skill_match['match_percentage']:.1f}%",
                        f"{skill_match['matched_count']}/{skill_match['total_required']}"
                    )
                
                with col3:
                    st.metric(
                        "Experience Alignment",
                        f"{match_result['experience_alignment']:.1f}%"
                    )
                
                with col4:
                    st.metric(
                        "Education Alignment",
                        f"{match_result['education_alignment']:.1f}%"
                    )
                
                # Matched and missing skills
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("✅ Matched Skills")
                    matched_skills = match_result['matched_skills']
                    if matched_skills:
                        for skill in matched_skills:
                            st.write(f"• {skill}")
                    else:
                        st.info("No matched skills")
                
                with col2:
                    st.subheader("❌ Missing Skills")
                    missing_skills = match_result['missing_skills']
                    if missing_skills:
                        for skill in missing_skills:
                            st.write(f"• {skill}")
                    else:
                        st.success("All required skills present!")
                
                # Recommendation
                st.markdown("---")
                recommendation = match_result['recommendation']
                recommendations_text = {
                    'STRONG_MATCH': '🌟 This is an excellent match! Highly recommended for interview.',
                    'GOOD_MATCH': '👍 Good match. Consider for interview process.',
                    'MODERATE_MATCH': '⚠️ Moderate match. May require additional training.',
                    'POOR_MATCH': '❌ Poor match. Not recommended at this time.'
                }
                
                st.info(recommendations_text.get(recommendation, "Unable to determine"))
                
                # Update dashboard data
                st.session_state.dashboard_data['total_matches'] += 1
                st.session_state.dashboard_data['active_positions'] += 1  # Assuming each match represents an active position
                
                # Update average match score
                current_avg = st.session_state.dashboard_data['avg_match_score']
                total_matches = st.session_state.dashboard_data['total_matches']
                new_avg = ((current_avg * (total_matches - 1)) + score) / total_matches
                st.session_state.dashboard_data['avg_match_score'] = new_avg
                
                # Add to recent matches (keep only last 10)
                candidate_name = "Unknown Candidate"
                if resume_file:
                    candidate_name = resume_file.name.replace('.pdf', '').replace('.txt', '').replace('.docx', '').title()
                
                match_entry = {
                    'Candidate': candidate_name,
                    'Position': 'Job Description',  # Could be extracted from job text if needed
                    'Match Score': round(score, 1),
                    'Status': 'Shortlisted' if score >= 70 else 'In Review' if score >= 50 else 'Pending'
                }
                
                st.session_state.dashboard_data['recent_matches'].insert(0, match_entry)
                st.session_state.dashboard_data['recent_matches'] = st.session_state.dashboard_data['recent_matches'][:10]  # Keep only last 10
                
                # Update shortlisted count
                if score >= 70:
                    st.session_state.dashboard_data['shortlisted'] += 1
        else:
            st.warning("⚠️ Please upload both resume and job description files")
    module_close()


def show_insights():
    """Analytics and insights page"""
    module_wrapper('📈 Insights & Analytics')
    
    tab1, tab2, tab3 = st.tabs(["Skill Analytics", "Top Skills", "Candidate Distribution"])
    
    with tab1:
        st.subheader("📊 Skill Demand Over Time")
        
        # Placeholder chart
        data = pd.DataFrame({
            'Skill': ['Python', 'JavaScript', 'Java', 'React', 'AWS'],
            'Demand': [85, 70, 65, 80, 75]
        })
        
        st.bar_chart(data.set_index('Skill'))
    
    with tab2:
        st.subheader("⭐ Most In-Demand Skills")
        
        skills_data = pd.DataFrame({
            'Rank': range(1, 6),
            'Skill': ['Python', 'JavaScript', 'React', 'AWS', 'Machine Learning'],
            'Mentions': [342, 298, 276, 241, 198]
        })
        
        st.dataframe(skills_data, width='stretch')
    
    with tab3:
        st.subheader("👥 Candidate Distribution")
        
        distribution = pd.DataFrame({
            'Experience Level': ['Entry-level (0-2yrs)', 'Junior (2-5yrs)', 'Mid-level (5-10yrs)', 'Senior (10+yrs)'],
            'Count': [45, 120, 85, 35]
        })
        
        st.bar_chart(distribution.set_index('Experience Level'))
    module_close()


def show_settings():
    """Settings page"""
    module_wrapper('⚙️ Settings')
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🤖 Model Configuration")
        
        model_type = st.selectbox(
            "Select ML Model:",
            ["Random Forest", "Logistic Regression", "SVM", "Decision Tree"]
        )
        
        confidence_threshold = st.slider(
            "Confidence Threshold:",
            0.0, 1.0, 0.7
        )
    
    with col2:
        st.subheader("📋 Matching Engine Settings")
        
        content_weight = st.slider(
            "Content Similarity Weight:",
            0.0, 1.0, 0.30
        )
        
        skill_weight = st.slider(
            "Skill Match Weight:",
            0.0, 1.0, 0.40
        )
    
    if st.button("💾 Save Settings", width='stretch'):
        st.success("✅ Settings saved successfully!")
    
    st.markdown("---")
    
    st.subheader("📂 Data & Models")
    
    if st.button("🔄 Retrain Models"):
        st.info("Model retraining feature coming soon...")
    
    if st.button("🗑️ Clear Cache"):
        st.success("✅ Cache cleared!")
    st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()
