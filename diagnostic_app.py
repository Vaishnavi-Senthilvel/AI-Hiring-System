"""
Diagnostic test for file upload functionality in Streamlit resume analysis
"""
import streamlit as st
import sys
from pathlib import Path
from io import BytesIO

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from nlp_processor import NLPProcessor

st.set_page_config(page_title="Resume Upload Diagnostic", layout="wide")

st.title("🔧 Resume Upload Diagnostic Tool")

st.markdown("---")
st.subheader("1️⃣ Test File Upload Widget")

uploaded_file = st.file_uploader(
    "Choose a resume file to test",
    type=['pdf', 'txt', 'docx'],
    help="Select any PDF, TXT, or DOCX file"
)

if uploaded_file:
    st.write("### File Details:")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("File Name", uploaded_file.name)
    with col2:
        st.metric("File Size (bytes)", uploaded_file.size)
    with col3:
        st.metric("MIME Type", uploaded_file.type or "Not detected")
    with col4:
        st.metric("File Format", uploaded_file.name.split('.')[-1].upper())
    
    st.write("---")
    st.write("### 2️⃣ Extract Text from File:")
    
    try:
        # Try to extract text
        file_name = uploaded_file.name.lower()
        
        # Handle text files
        if uploaded_file.type == "text/plain" or file_name.endswith('.txt'):
            st.info("Detected: TXT file")
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            st.success(f"✅ Successfully extracted {len(content)} characters")
            st.text_area("Extracted Content", content, height=200)
        
        # Handle PDF files
        elif uploaded_file.type == "application/pdf" or file_name.endswith('.pdf'):
            st.info("Detected: PDF file")
            try:
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(BytesIO(uploaded_file.read()))
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                
                if text.strip():
                    st.success(f"✅ Successfully extracted {len(text)} characters from {len(pdf_reader.pages)} pages")
                    st.text_area("Extracted Content", text, height=200)
                else:
                    st.warning("⚠️ PDF extracted but no text content found (may be scanned image)")
            except ImportError:
                st.error("❌ PyPDF2 library not installed")
            except Exception as e:
                st.error(f"❌ Error reading PDF: {str(e)}")
        
        # Handle DOCX files
        elif file_name.endswith(('.docx', '.doc')):
            st.info("Detected: DOCX/DOC file")
            try:
                from docx import Document
                doc = Document(BytesIO(uploaded_file.read()))
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"
                
                if text.strip():
                    st.success(f"✅ Successfully extracted {len(text)} characters")
                    st.text_area("Extracted Content", text, height=200)
                else:
                    st.warning("⚠️ DOCX extracted but no text content found")
            except ImportError:
                st.error("❌ python-docx library not installed")
            except Exception as e:
                st.error(f"❌ Error reading DOCX: {str(e)}")
        else:
            st.error(f"❌ File type not supported: {uploaded_file.type}")
    
    except Exception as e:
        st.error(f"❌ Unexpected error: {str(e)}")
        import traceback
        st.text(traceback.format_exc())
    
    st.write("---")
    st.write("### 3️⃣ Test NLP Processing:")
    
    try:
        nlp = NLPProcessor()
        st.success("✅ NLPProcessor initialized")
        
        # Extract information
        skills = nlp.extract_skills_from_text(content) if 'content' in locals() else []
        education = nlp.extract_education(content) if 'content' in locals() else []
        experience = nlp.extract_years_of_experience(content) if 'content' in locals() else 0
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Skills Found", len(skills))
            if skills:
                st.write(", ".join(skills[:5]))
        
        with col2:
            st.metric("Education Found", len(education))
            if education:
                st.write(", ".join(education[:3]))
        
        with col3:
            st.metric("Experience (years)", experience)
    
    except Exception as e:
        st.error(f"❌ NLP processing error: {str(e)}")
        import traceback
        st.text(traceback.format_exc())

else:
    st.info("👆 Upload a resume file above to start testing")

st.markdown("---")
st.subheader("💡 Troubleshooting Tips:")
st.write("""
- ✅ Ensure file is not password-protected
- ✅ Try uploading a TXT file first to test basic functionality
- ✅ Make sure the file is not empty
- ✅ Check that PyPDF2 and python-docx are installed
- ✅ For PDFs, ensure they're not scanned images (text-based PDFs work best)
""")
