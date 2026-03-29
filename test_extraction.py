#!/usr/bin/env python3
"""
Test script for text extraction functionality
"""

from io import BytesIO

def extract_text_from_file(uploaded_file):
    """Extract text content from uploaded file (PDF, TXT, DOCX)"""
    try:
        file_name = uploaded_file.name.lower()

        # Handle text files
        if uploaded_file.type == "text/plain" or file_name.endswith('.txt'):
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            uploaded_file.seek(0)  # Reset file pointer
            return content

        # Handle PDF files
        elif uploaded_file.type == "application/pdf" or file_name.endswith('.pdf'):
            try:
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(BytesIO(uploaded_file.read()))
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text += page_text + "\n"
                return text.strip()
            except ImportError:
                return "PDF processing library not installed. Please install PyPDF2: pip install PyPDF2"
            except Exception as e:
                return f"Error processing PDF: {str(e)}. Try converting to text format."

        # Handle DOCX files
        elif (uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              or file_name.endswith(('.docx', '.doc'))):
            try:
                from docx import Document
                doc = Document(BytesIO(uploaded_file.read()))
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        text += paragraph.text + "\n"
                # Also extract from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"
                return text.strip()
            except ImportError:
                return "DOCX processing library not installed. Please install python-docx: pip install python-docx"
            except Exception as e:
                return f"Error processing DOCX: {str(e)}. Try converting to text format."

        else:
            return f"Unsupported file type: {uploaded_file.type}. Please upload PDF, TXT, or DOCX files."

    except Exception as e:
        return f"Error extracting text from {uploaded_file.name}: {str(e)}"


def test_text_extraction():
    """Test the text extraction function with sample data"""
    print("Testing text extraction functionality...")

    # Test imports
    try:
        from PyPDF2 import PdfReader
        print("✅ PyPDF2 imported successfully")
    except ImportError:
        print("❌ PyPDF2 not available")

    try:
        from docx import Document
        print("✅ python-docx imported successfully")
    except ImportError:
        print("❌ python-docx not available")

    # Test basic functionality
    print("\nTesting basic text processing...")

    # Test TXT extraction
    txt_content = "This is a test resume.\nIt contains skills like Python, JavaScript."

    class MockFile:
        def __init__(self, content, file_type, file_name):
            self.type = file_type
            self.name = file_name
            self._content = content
            self._position = 0

        def read(self):
            if self._position == 0:
                self._position = len(self._content)
                return self._content
            return b''

        def seek(self, position):
            self._position = position

    txt_file_mock = MockFile(txt_content.encode('utf-8'), 'text/plain', 'test.txt')

    result = extract_text_from_file(txt_file_mock)
    print(f"TXT extraction result: '{result}'")
    print(f"Length: {len(result)} characters")

    # Test with empty content
    empty_file_mock = MockFile(b'', 'text/plain', 'empty.txt')

    result_empty = extract_text_from_file(empty_file_mock)
    print(f"Empty file result: '{result_empty}'")

    print("✅ Text extraction test completed")

if __name__ == "__main__":
    test_text_extraction()